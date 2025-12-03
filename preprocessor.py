import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
from datetime import datetime
import math
import os
import tempfile
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
except ImportError: 
    Document = None
    Inches = None

try:
    import matplotlib.pyplot as plt
except ImportError:  
    plt = None

PASTEL_BG = "#e6f0ff"
PASTEL_BUTTON = "#99c2ff"
PASTEL_HOVER = "#b3d1ff"
PASTEL_TREE_BG = "#cce0ff"
PASTEL_TREE_ALT = "#e6f2ff"
PASTEL_TREE_BORDER = "#4da6ff"

DEFAULT_A = 1.0
DEFAULT_E = 1.0
DEFAULT_SIGMA = 1.0

#Оформление
def create_rounded_rect(canvas, x1, y1, x2, y2, radius=10, **kwargs):
    points = [
        x1+radius, y1,
        x2-radius, y1,
        x2, y1, x2, y1+radius,
        x2, y2-radius,
        x2, y2, x2-radius, y2,
        x1+radius, y2,
        x1, y2, x1, y2-radius,
        x1, y1+radius,
        x1, y1
    ]
    return canvas.create_polygon(points, smooth=True, **kwargs)

class RoundedButton(tk.Frame):
    def __init__(self, master, text, command=None, width=150, height=28):
        super().__init__(master, bg=PASTEL_BG)
        self.command = command
        self.canvas = tk.Canvas(self, width=width, height=height, highlightthickness=0, bg=PASTEL_BG)
        self.canvas.pack()
        self.rect = create_rounded_rect(self.canvas, 2, 2, width-2, height-2, radius=8,
                                        outline="#4da6ff", width=2, fill=PASTEL_BUTTON)
        self.text = self.canvas.create_text(width/2, height/2, text=text, font=("Arial", 10, "bold"))
        self.canvas.bind("<Button-1>", lambda e: self.on_click())
        self.canvas.bind("<Enter>", lambda e: self.canvas.itemconfig(self.rect, fill=PASTEL_HOVER))
        self.canvas.bind("<Leave>", lambda e: self.canvas.itemconfig(self.rect, fill=PASTEL_BUTTON))

    def on_click(self):
        if self.command:
            self.command()


class StructuralProcessor:
    """Вычисляет перемещения и внутренние усилия плоской стержневой системы."""

    def __init__(self, nodes, elements, supports, forces, element_forces):
        self.nodes = nodes
        self.elements = elements
        self.supports = supports
        self.forces = forces
        self.element_forces = element_forces

    def solve(self):
        self._validate_input()
        dof = len(self.nodes)
        K = np.zeros((dof, dof), dtype=float)
        F = np.zeros(dof, dtype=float)

        # Сборка матрицы жесткости
        for idx, element in enumerate(self.elements):
            n1, n2 = element['nodes']
            L = element['L']
            stiffness = element['A'] * element['E'] / L
            k_local = stiffness * np.array([[1, -1], [-1, 1]], dtype=float)
            dof_indices = [n1 - 1, n2 - 1]
            for i in range(2):
                for j in range(2):
                    K[dof_indices[i], dof_indices[j]] += k_local[i, j]

        # Сосредоточенные силы
        for node_idx, value in self.forces:
            F[node_idx - 1] += value

        # Эквивалентные нагрузки от q
        q_map = {idx: q for idx, q in self.element_forces}
        for idx, q_val in q_map.items():
            element = self.elements[idx]
            L = element['L']
            eq = q_val * L / 2
            n1, n2 = element['nodes']
            F[n1 - 1] += eq
            F[n2 - 1] += eq

        K_original = K.copy()
        F_original = F.copy()

        K_bc, F_bc = self._apply_supports(K, F)

        try:
            displacements = np.linalg.solve(K_bc, F_bc)
        except np.linalg.LinAlgError as exc:
            raise ValueError("Матрица системы вырождена. Проверьте граничные условия.") from exc

        reactions = (K_original @ displacements) - F_original
        element_results = self._compute_element_results(displacements, q_map)
        warnings = self._build_warnings(element_results)

        return {
            'displacements': displacements.tolist(),
            'element_results': element_results,
            'reactions': [(node, float(reactions[node - 1])) for node in self.supports],
            'warnings': warnings,
        }

    def _apply_supports(self, stiffness, loads):
        if not self.supports:
            raise ValueError("Добавьте хотя бы одну опору для расчёта.")

        K = stiffness.copy()
        F = loads.copy()
        for node in self.supports:
            idx = node - 1
            K[idx, :] = 0.0
            K[:, idx] = 0.0
            K[idx, idx] = 1.0
            F[idx] = 0.0
        return K, F

    def _compute_element_results(self, displacements, q_map):
        results = []
        for idx, element in enumerate(self.elements):
            n1, n2 = element['nodes']
            u1 = displacements[n1 - 1]
            u2 = displacements[n2 - 1]
            L = element['L']

            stress = element['E'] * ((u2 - u1) / L) 

            axial_force = stress * element['A']
            q_val = q_map.get(idx, 0.0)

            results.append({
                'index': idx + 1,
                'nodes': (n1, n2),
                'length': L,
                'u1': float(u1),
                'u2': float(u2),
                'stress': float(stress),
                'force': float(axial_force),
                'allowable_sigma': element['sigma'],
                'A': element['A'],
                'E': element['E'],
                'q': float(q_val),
            })
        return results


    def _build_warnings(self, element_results):
        warnings = []
        for res in element_results:
            if abs(res['stress']) > res['allowable_sigma']:
                warnings.append(
                    f"Стержень {res['index']} (узлы {res['nodes'][0]}-{res['nodes'][1]}) превышает допускаемое напряжение."
                )
        return warnings

    def _validate_input(self):
        if not self.nodes or len(self.nodes) < 2:
            raise ValueError("Добавьте минимум два узла.")
        if not self.elements:
            raise ValueError("Добавьте минимум один стержень.")

    @staticmethod
    def evaluate_section(element_result, x_local):
        L = element_result['length']
        if not (0.0 <= x_local <= L):
            raise ValueError("Локальная координата должна быть в пределах [0, L].")

        xi = x_local / L
        u = element_result['u1'] * (1 - xi) + element_result['u2'] * xi
        if element_result['q'] != 0:

            E = element_result['E']
            A = element_result['A']
            parabolic_term = (element_result['q'] * L * L / (2 * E * A)) * xi * (1 - xi)
            u += parabolic_term
        if element_result['q'] == 0:
            nx = element_result['force']
        else:
            nx = element_result['force'] - element_result['q'] * (x_local - L / 2)
        sigma = nx / element_result['A']
        return {'ux': u, 'Nx': nx, 'sigma': sigma}


class DiagramWindow(tk.Toplevel):
    def __init__(self, master, segments, title, y_label, component):
        super().__init__(master)
        self.title(title)
        self.geometry("800x500")
        self.configure(bg=PASTEL_BG)
        self.segments = segments
        self.y_label = y_label
        self.component = component

        # Верхняя часть: холст с эпюрой
        self.canvas = tk.Canvas(self, bg="white", height=320)
        self.canvas.pack(fill=tk.BOTH, expand=False)
        self.canvas.bind("<Configure>", lambda e: self._draw())

        # Нижняя часть: таблица значений по точкам
        table_frame = ttk.Frame(self)
        table_frame.pack(fill=tk.BOTH, expand=True)

        columns = ("elem", "x", "ux", "Nx", "sigma")
        self.points_table = ttk.Treeview(table_frame, columns=columns, show="headings", height=6)
        self.points_table.heading("elem", text="Стержень")
        self.points_table.heading("x", text="x, м")
        self.points_table.heading("ux", text="ux, м")
        self.points_table.heading("Nx", text="Nx, Н")
        self.points_table.heading("sigma", text="σx, Па")

        for col in columns:
            self.points_table.column(col, width=100, anchor=tk.CENTER)

        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.points_table.yview)
        self.points_table.configure(yscrollcommand=vsb.set)
        self.points_table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        self._fill_points_table()

    def _fill_points_table(self):
        """Заполнить таблицу значений по всем точкам эпюры."""
        for row in self.points_table.get_children():
            self.points_table.delete(row)

        for seg in self.segments:
            elem_idx = seg.get('index', '')
            for val in seg.get('values', []):
                self.points_table.insert(
                    "",
                    tk.END,
                    values=(
                        f"{elem_idx}",
                        f"{val['x']:.4f}",
                        f"{val['ux']:.6f}",
                        f"{val['Nx']:.2f}",
                        f"{val['sigma']:.2f}",
                    )
                )

    @staticmethod
    def _format_value(value):
        abs_val = abs(value)
        if abs_val == 0:
            return "0"
        if abs_val >= 1:
            return f"{value:.2f}"
        if abs_val >= 1e-2:
            return f"{value:.4f}"
        return f"{value:.2e}"

    def _draw_point_label(self, x, y, value, plot_top, plot_bottom, prefix=""):
        offset = -12 if value >= 0 else 12
        y_label = max(plot_top + 12, min(plot_bottom - 12, y + offset))
        label_text = f"{prefix}{self._format_value(value)}"
        self.canvas.create_text(
            x,
            y_label,
            text=label_text,
            font=("Arial", 8, "bold"),
            fill="#08306b"
        )

    def _label_extrema_points(self, points, plot_top, plot_bottom):
        if self.component != 'ux' or len(points) < 2:
            return
        tol = 1e-10
        extrema_points = []

        for i in range(1, len(points) - 1):
            prev_val = points[i - 1]['value']
            curr_val = points[i]['value']
            next_val = points[i + 1]['value']
            delta_prev = curr_val - prev_val
            delta_next = next_val - curr_val
            if abs(delta_prev) < tol and abs(delta_next) < tol:
                continue
            if delta_prev > 0 and delta_next < 0:
                extrema_points.append(points[i])
            elif delta_prev < 0 and delta_next > 0:
                extrema_points.append(points[i])

        max_val = max(points, key=lambda p: p['value'])['value']
        min_val = min(points, key=lambda p: p['value'])['value']
        if max_val - min_val > tol:
            if math.isclose(points[0]['value'], max_val, rel_tol=1e-6, abs_tol=tol):
                extrema_points.append(points[0])
            if math.isclose(points[-1]['value'], max_val, rel_tol=1e-6, abs_tol=tol):
                extrema_points.append(points[-1])
            if math.isclose(points[0]['value'], min_val, rel_tol=1e-6, abs_tol=tol):
                extrema_points.append(points[0])
            if math.isclose(points[-1]['value'], min_val, rel_tol=1e-6, abs_tol=tol):
                extrema_points.append(points[-1])

        seen = set()
        for point in extrema_points:
            key = (round(point['x'], 2), round(point['value'], 6))
            if key in seen:
                continue
            seen.add(key)
            self.canvas.create_oval(
                point['x'] - 4,
                point['y'] - 4,
                point['x'] + 4,
                point['y'] + 4,
                outline="#d62728",
                width=2
            )
            self._draw_point_label(point['x'], point['y'], point['value'], plot_top, plot_bottom)

    def _draw(self):
        self.canvas.delete("all")
        if not self.segments:
            return

        width = self.canvas.winfo_width()
        height = self.canvas.winfo_height()
        margin = 60

        values = [val for seg in self.segments for _, val in seg['points']]
        min_val = min(values)
        max_val = max(values)
        if math.isclose(max_val, min_val):
            max_val += 1.0
            min_val -= 1.0
        scale_y = (height - 2 * margin) / (max_val - min_val)

        plot_left = margin
        plot_right = width - margin
        plot_top = margin
        plot_bottom = height - margin

        zero_y = plot_bottom - (0 - min_val) * scale_y

        # Сетка
        num_x_grid = 10
        for i in range(num_x_grid + 1):
            x_grid = plot_left + (plot_right - plot_left) * i / num_x_grid
            self.canvas.create_line(x_grid, plot_top, x_grid, plot_bottom, fill="#e0e0e0", width=1)
        num_y_grid = 8
        for i in range(num_y_grid + 1):
            y_grid = plot_top + (plot_bottom - plot_top) * i / num_y_grid
            val_grid = max_val - (max_val - min_val) * i / num_y_grid
            self.canvas.create_line(plot_left, y_grid, plot_right, y_grid, fill="#e0e0e0", width=1)
            self.canvas.create_text(plot_left - 5, y_grid, text=f"{val_grid:.2f}", font=("Arial", 7), fill="#666666", anchor=tk.E)

        # Оси
        self.canvas.create_rectangle(plot_left, plot_top, plot_right, plot_bottom, outline="#b3cde4")
        self.canvas.create_line(plot_left, zero_y, plot_right, zero_y, fill="#888888", dash=(4, 2), width=2)
        self.canvas.create_line(plot_left, plot_bottom, plot_right, plot_bottom, arrow=tk.LAST, width=2)
        self.canvas.create_line(plot_left, plot_bottom, plot_left, plot_top, arrow=tk.LAST, width=2)
        self.canvas.create_text(plot_left - 25, plot_top + 10, text=self.y_label, angle=90, font=("Arial", 10, "bold"))
        self.canvas.create_text((plot_left + plot_right) / 2, plot_bottom + 25, text="Локальная координата x, м", font=("Arial", 10, "bold"))

        plot_width = plot_right - plot_left
        segment_count = len(self.segments)
        total_length = sum(seg['length'] for seg in self.segments)
        if total_length <= 0:
            total_length = 1.0
        scale_x = plot_width / total_length

        current_x = plot_left
        # Палитра цветов для разных стержней
        segment_colors = [
            ("#1f77b4", "#aec7e8"),
            ("#ff7f0e", "#ffbb78"),
            ("#2ca02c", "#98df8a"),
            ("#d62728", "#ff9896"),
            ("#9467bd", "#c5b0d5"),
            ("#8c564b", "#c49c94"),
            ("#e377c2", "#f7b6d2"),
            ("#7f7f7f", "#c7c7c7"),
            ("#bcbd22", "#dbdb8d"),
            ("#17becf", "#9edae5"),
        ]
        all_points = []
        for idx_seg, seg in enumerate(self.segments):
            line_color, fill_color = segment_colors[idx_seg % len(segment_colors)]
            pts = []
            for x_local, value in seg['points']:
                x = current_x + x_local * scale_x
                y = height - margin - (value - min_val) * scale_y
                pts.append((x, y, value))
                all_points.append({'x': x, 'y': y, 'value': value})
            line_pts = [(x, y) for x, y, _ in pts]
            if len(line_pts) >= 2:
                # Линия идёт по тем же точкам, что и заливка (без сглаживания),
                # чтобы не было расхождения между кривой и заливкой при малом числе точек.
                self.canvas.create_line(line_pts, fill=line_color, width=3)
                for i in range(len(line_pts) - 1):
                    x1, y1 = line_pts[i]
                    x2, y2 = line_pts[i + 1]
                    self.canvas.create_polygon(
                        x1, zero_y, x1, y1, x2, y2, x2, zero_y,
                        fill=fill_color, outline="", stipple="gray50"
                    )

            if pts:
                x, y, value = pts[0]
                self.canvas.create_oval(x - 3, y - 3, x + 3, y + 3, fill=line_color, outline="")
                self._draw_point_label(x, y, value, plot_top, plot_bottom)

                if len(pts) > 1:
                    x, y, value = pts[-1]
                    self.canvas.create_oval(x - 3, y - 3, x + 3, y + 3, fill=line_color, outline="")
                    self._draw_point_label(x, y, value, plot_top, plot_bottom)
            label_x = (line_pts[0][0] + line_pts[-1][0]) / 2 if line_pts else current_x
            self.canvas.create_text(label_x, plot_bottom + 12, text=f"e{seg['index']}", font=("Arial", 9, "bold"))
            current_x += seg['length'] * scale_x
        self._label_extrema_points(all_points, plot_top, plot_bottom)

class PreprocessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Плоская стержневая система")
        self.root.geometry("1600x800")
        self.root.configure(bg=PASTEL_BG)

        self.nodes = []             
        self.elements = []  
        self.supports = []   
        self.forces = [] 
        self.element_forces = []
        self.view_min = None
        self.view_max = None
        self.results = None
        self.section_result_var = tk.StringVar(value="Результаты отсутствуют")
        self.sorted_elements_for_combo = []
        self.imported_file_name = None 

        self.create_ui()
        self.draw_system()

    #Окно
    def create_ui(self):
        control_container = ttk.Frame(self.root)
        control_container.pack(side=tk.LEFT, fill=tk.Y)

        canvas_ctrl = tk.Canvas(control_container, width=280, bg=PASTEL_BG, highlightthickness=0)
        canvas_ctrl.pack(side=tk.LEFT, fill=tk.Y, expand=True)

        scrollbar = ttk.Scrollbar(control_container, orient="vertical", command=canvas_ctrl.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas_ctrl.configure(yscrollcommand=scrollbar.set)

        self.control_frame = tk.Frame(canvas_ctrl, bg=PASTEL_BG)
        canvas_ctrl.create_window((0,0), window=self.control_frame, anchor="nw")
        self.control_frame.bind("<Configure>", lambda e: canvas_ctrl.configure(scrollregion=canvas_ctrl.bbox("all")))

        tk.Label(self.control_frame, text="Добавление узла", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=5)
        tk.Label(self.control_frame, text="Координата X:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_node = ttk.Entry(self.control_frame,width=15)
        self.entry_node.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Добавить узел", self.add_node).pack(fill=tk.X, pady=2)
        RoundedButton(self.control_frame, "Удалить узел", self.delete_node).pack(fill=tk.X, pady=2)

        tk.Label(self.control_frame, text="Добавление стержня", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=5)
        tk.Label(self.control_frame, text="Начальный узел n1:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.combo_n1 = ttk.Combobox(self.control_frame, width=5, state='readonly')
        self.combo_n1.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Конечный узел n2:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.combo_n2 = ttk.Combobox(self.control_frame, width=5, state='readonly')
        self.combo_n2.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Площадь A:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_A = ttk.Entry(self.control_frame,width=10)
        self.entry_A.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Модуль упругости E:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_E = ttk.Entry(self.control_frame,width=10)
        self.entry_E.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Допускаемое напряжение σ:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_sigma = ttk.Entry(self.control_frame,width=10)
        self.entry_sigma.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Добавить стержень", self.add_element).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Удалить стержень", self.delete_element).pack(fill=tk.X,pady=2)

        tk.Label(self.control_frame, text="Распределённая нагрузка q", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=5)
        tk.Label(self.control_frame, text="Выберите стержень:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.combo_q_element = ttk.Combobox(self.control_frame, width=10, state='readonly')
        self.combo_q_element.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Значение q:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_q = ttk.Entry(self.control_frame,width=10)
        self.entry_q.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Добавить нагрузку q", self.add_q).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Удалить нагрузку q", self.delete_q).pack(fill=tk.X,pady=2)

        tk.Label(self.control_frame, text="Добавление опоры", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=5)
        tk.Label(self.control_frame, text="Выберите узел:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.combo_support = ttk.Combobox(self.control_frame,width=5,state='readonly')
        self.combo_support.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Добавить опору", self.add_support).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Удалить опору", self.delete_support).pack(fill=tk.X,pady=2)

        tk.Label(self.control_frame, text="Силы F на узлах", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=5)
        tk.Label(self.control_frame, text="Выберите узел:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.combo_force_node = ttk.Combobox(self.control_frame,width=5,state='readonly')
        self.combo_force_node.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Значение F:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_F = ttk.Entry(self.control_frame,width=10)
        self.entry_F.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Добавить силу F", self.add_force).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Удалить силу F", self.delete_force).pack(fill=tk.X,pady=2)

        # Метка для отображения имени импортированного файла
        self.file_name_label = tk.Label(self.control_frame, text="", bg=PASTEL_BG, font=("Arial", 9), fg="#666666", wraplength=240)
        self.file_name_label.pack(pady=5)
        
        tk.Label(self.control_frame, text="Управление проектом", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=10)
        RoundedButton(self.control_frame, "Очистить всё", self.clear_all).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Масштаб +", lambda:self.scale_canvas(1.2)).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Масштаб -", lambda:self.scale_canvas(0.8)).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Сброс масштаба", self.reset_scale).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Экспорт JSON", self.export_json).pack(fill=tk.X,pady=2)
        RoundedButton(self.control_frame, "Импорт JSON", self.import_json).pack(fill=tk.X,pady=2)

        tk.Label(self.control_frame, text="Расчёт", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=10)
        RoundedButton(self.control_frame, "Рассчитать систему", self.run_analysis).pack(fill=tk.X, pady=2)
        RoundedButton(self.control_frame, "Экспорт отчёта", self.export_report).pack(fill=tk.X, pady=2)

        tk.Label(self.control_frame, text="Компоненты в сечении", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=10)
        tk.Label(self.control_frame, text="Выберите стержень:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.section_element_combo = ttk.Combobox(self.control_frame, width=10, state='readonly')
        self.section_element_combo.pack(anchor=tk.W)
        tk.Label(self.control_frame, text="Локальная координата x:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_section_coordinate = ttk.Entry(self.control_frame, width=10)
        self.entry_section_coordinate.pack(anchor=tk.W)
        RoundedButton(self.control_frame, "Показать компоненты", self.calculate_section_values).pack(fill=tk.X, pady=2)
        tk.Label(self.control_frame, textvariable=self.section_result_var, wraplength=240, bg=PASTEL_BG, justify=tk.LEFT).pack(fill=tk.X, pady=4)

        tk.Label(self.control_frame, text="Диаграммы", font=("Arial",11,"bold"), bg=PASTEL_BG).pack(pady=10)
        tk.Label(self.control_frame, text="Точек на стержень:", bg=PASTEL_BG).pack(anchor=tk.W)
        self.entry_diagram_points = ttk.Entry(self.control_frame, width=10)
        self.entry_diagram_points.insert(0, "20")
        self.entry_diagram_points.pack(anchor=tk.W, pady=(0, 4))
        RoundedButton(self.control_frame, "Эпюра ux", lambda: self.show_diagram('ux')).pack(fill=tk.X, pady=2)
        RoundedButton(self.control_frame, "Эпюра Nx", lambda: self.show_diagram('Nx')).pack(fill=tk.X, pady=2)
        RoundedButton(self.control_frame, "Эпюра σx", lambda: self.show_diagram('sigma')).pack(fill=tk.X, pady=2)

        right_frame = ttk.Frame(self.root)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        canvas_container = ttk.Frame(right_frame)
        canvas_container.pack(fill=tk.BOTH, expand=True)
        canvas_container.columnconfigure(0, weight=1)
        canvas_container.rowconfigure(0, weight=1)

        self.canvas = tk.Canvas(canvas_container, bg="white")
        self.canvas.grid(row=0, column=0, sticky="nsew")
        canvas_vscroll = ttk.Scrollbar(canvas_container, orient=tk.VERTICAL, command=self.canvas.yview)
        canvas_hscroll = ttk.Scrollbar(canvas_container, orient=tk.HORIZONTAL, command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=canvas_vscroll.set, xscrollcommand=canvas_hscroll.set)
        canvas_vscroll.grid(row=0, column=1, sticky="ns")
        canvas_hscroll.grid(row=1, column=0, sticky="ew")

        self.canvas.bind("<MouseWheel>", self.on_mousewheel)
        self.canvas.bind("<Button-4>", lambda e: self.scale_canvas(1.1, e.x))
        self.canvas.bind("<Button-5>", lambda e: self.scale_canvas(0.9, e.x))

        tabs_container = ttk.Frame(right_frame)
        tabs_container.pack(fill=tk.BOTH, expand=True)
        self.tab_control = ttk.Notebook(tabs_container)
        self.tab_control.pack(fill=tk.BOTH, expand=True)

        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Treeview",
                        background=PASTEL_TREE_BG,
                        fieldbackground=PASTEL_TREE_BG,
                        foreground="black",
                        rowheight=25,
                        font=("Arial", 10))
        style.map("Treeview", background=[('selected', '#66b3ff')])
        style.configure("Treeview.Heading", font=("Arial", 10, "bold"), background=PASTEL_BUTTON, foreground="black")

        #Вкладки с табличками
        self.tree_nodes = self._create_tree_tab("Узлы", ("X",), ("X, м",))
        self.tree_elements = self._create_tree_tab(
            "Стержни",
            ("nodes", "L", "A", "E", "sigma"),
            ("Стержень", "L, м", "A, м²", "E, Па", "σдоп, Па")
        )
        self.tree_q = self._create_tree_tab(
            "Нагрузки q",
            ("nodes", "q"),
            ("Стержень", "q, Н/м")
        )
        self.tree_supports = self._create_tree_tab("Опоры", ("node",), ("Узел",))
        self.tree_forces = self._create_tree_tab(
            "Силы F",
            ("node", "F"),
            ("Узел", "F, Н")
        )
        self.tree_displacements = self._create_tree_tab(
            "Перемещения",
            ("node", "ux"),
            ("Узел", "ux, м")
        )
        self.tree_element_results = self._create_tree_tab(
            "Результаты расчетов",
            ("element","Nx1","Nx2","sigma1","sigma2","u1","u2"),
            ("Стержень","Nx1, Н","Nx2, Н","σx1, Па","σx2, Па","u1, м","u2, м")
        )

        analysis_frame = ttk.Frame(self.tab_control)
        self.analysis_text = tk.Text(analysis_frame, height=10, wrap="word", state='disabled', bg="white")
        self.analysis_text.pack(fill=tk.BOTH, expand=True)
        self.tab_control.add(analysis_frame, text="Анализ")

        self.reset_postprocessor()

    def _create_tree_tab(self, title, columns, headings):
        frame = ttk.Frame(self.tab_control)
        
        tree_container = ttk.Frame(frame)
        tree_container.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        tree = ttk.Treeview(tree_container, columns=columns, show='headings')
        for col, head in zip(columns, headings):
            tree.heading(col, text=head)
            tree.column(col, width=100, minwidth=80) 
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        

        v_scrollbar = ttk.Scrollbar(tree_container, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=v_scrollbar.set)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tab_control.add(frame, text=title)
        return tree

    #Обновление окошек с выбором
    def update_comboboxes(self):
        node_str = [str(i+1) for i in range(len(self.nodes))] if self.nodes else []
        self.combo_n1['values'] = node_str
        self.combo_n2['values'] = node_str
        self.combo_support['values'] = node_str
        self.combo_force_node['values'] = node_str

        if self.elements:
            sorted_with_index = sorted(enumerate(self.elements), key=lambda x: min(self.nodes[x[1]['nodes'][0]-1], self.nodes[x[1]['nodes'][1]-1]))
            self.sorted_elements_for_combo = sorted_with_index
            element_str = [f"{e['nodes'][0]}-{e['nodes'][1]}" for idx,e in sorted_with_index]
        else:
            self.sorted_elements_for_combo = []
            element_str = []

        self.combo_q_element['values'] = element_str
        if hasattr(self, 'section_element_combo'):
            self.section_element_combo['values'] = element_str

        self.combo_n1.set('')
        self.combo_n2.set('')
        self.combo_q_element.set('')
        self.combo_support.set('')
        self.combo_force_node.set('')


    #Обновление табличек
    def refresh_lists(self):
            self.tree_nodes.delete(*self.tree_nodes.get_children())
            for i,x in enumerate(sorted(self.nodes), start=1):
                self.tree_nodes.insert("", tk.END, values=(f"{x:.2f}",))

            self.tree_elements.delete(*self.tree_elements.get_children())
            sorted_elements = sorted(self.elements, key=lambda e: min(self.nodes[e['nodes'][0]-1], self.nodes[e['nodes'][1]-1]))
            for e in sorted_elements:
                n1,n2 = e['nodes']
                self.tree_elements.insert("", tk.END, values=(f"{n1}-{n2}", f"{e['L']:.2f}", f"{e['A']:.2f}", f"{e['E']:.2e}", f"{e['sigma']:.2e}"))

            self.tree_q.delete(*self.tree_q.get_children())
            sorted_q = []
            for idx, q_val in self.element_forces:
                e = self.elements[idx]
                sorted_q.append((min(self.nodes[e['nodes'][0]-1], self.nodes[e['nodes'][1]-1]), idx, q_val))
            sorted_q.sort()
            for _, idx, q_val in sorted_q:
                e = self.elements[idx]
                n1,n2 = e['nodes']
                self.tree_q.insert("", tk.END, values=(f"{n1}-{n2}", f"{q_val:.2f}"))

            self.tree_supports.delete(*self.tree_supports.get_children())
            for n in sorted(self.supports):
                self.tree_supports.insert("", tk.END, values=(n,))

            self.tree_forces.delete(*self.tree_forces.get_children())
            for n,F in sorted(self.forces, key=lambda x:x[0]):
                self.tree_forces.insert("", tk.END, values=(n,F))

    def reset_postprocessor(self):
        self.results = None
        if hasattr(self, 'tree_displacements'):
            self.tree_displacements.delete(*self.tree_displacements.get_children())
        if hasattr(self, 'tree_element_results'):
            self.tree_element_results.delete(*self.tree_element_results.get_children())
        if hasattr(self, 'analysis_text'):
            self.analysis_text.configure(state='normal')
            self.analysis_text.delete("1.0", tk.END)
            self.analysis_text.insert(tk.END, "Результаты отсутствуют.\n")
            self.analysis_text.configure(state='disabled')
        if hasattr(self, 'section_element_combo'):
            self.section_element_combo.set('')
        self.section_result_var.set("Результаты отсутствуют")

    def populate_results_ui(self):
        if not self.results:
            return
        self.tree_displacements.delete(*self.tree_displacements.get_children())
        for i, ux in enumerate(self.results['displacements'], start=1):
            self.tree_displacements.insert("", tk.END, values=(i, f"{ux:.5f}"))

        self.tree_element_results.delete(*self.tree_element_results.get_children())
        for res in self.results['element_results']:
            length = res['length']
            q_val = res['q']

            Nx1 = res['force'] + q_val * length / 2
            Nx2 = res['force'] - q_val * length / 2

            sigma1 = Nx1 / res['A']
            sigma2 = Nx2 / res['A']
            self.tree_element_results.insert(
                "",
                tk.END,
                values=(
                    f"{res['nodes'][0]}-{res['nodes'][1]}",
                    f"{Nx1:.2f}",
                    f"{Nx2:.2f}",
                    f"{sigma1:.2f}",
                    f"{sigma2:.2f}",
                    f"{res['u1']:.5f}",
                    f"{res['u2']:.5f}",
                )
            )

        self.analysis_text.configure(state='normal')
        self.analysis_text.delete("1.0", tk.END)
        
        self.analysis_text.insert(tk.END, "АНАЛИЗ РЕЗУЛЬТАТОВ РАСЧЁТА\n")
        
        self.analysis_text.insert(tk.END, "1. РЕАКЦИИ ОПОР:\n")
        total_reaction = 0.0
        for node, value in self.results.get('reactions', []):
            self.analysis_text.insert(tk.END, f"   R{node} = {value:.2f} Н\n")
            total_reaction += value
        if self.results.get('reactions'):
            self.analysis_text.insert(tk.END, f"   Сумма реакций: {total_reaction:.2f} Н\n")
        self.analysis_text.insert(tk.END, "\n")
        
        self.analysis_text.insert(tk.END, "2. ПЕРЕМЕЩЕНИЯ УЗЛОВ:\n")
        max_disp = max(abs(u) for u in self.results['displacements']) if self.results['displacements'] else 0
        max_disp_node = None
        for i, ux in enumerate(self.results['displacements'], start=1):
            self.analysis_text.insert(tk.END, f"   Узел {i}: ux = {ux:.6f} м\n")
            if abs(ux) == max_disp:
                max_disp_node = i
        if max_disp_node:
            self.analysis_text.insert(tk.END, f"   Максимальное перемещение: узел {max_disp_node}, |ux| = {max_disp:.6f} м\n")
        self.analysis_text.insert(tk.END, "\n")
        
        self.analysis_text.insert(tk.END, "3. ВНУТРЕННИЕ УСИЛИЯ И НАПРЯЖЕНИЯ:\n")
        max_force = 0.0
        max_force_elem = None
        max_stress = 0.0
        max_stress_elem = None
        for res in self.results['element_results']:
            length = res['length']
            q_val = res['q']
            Nx1 = res['force'] + q_val * length / 2
            Nx2 = res['force'] - q_val * length / 2
            sigma1 = abs(Nx1 / res['A'])
            sigma2 = abs(Nx2 / res['A'])
            max_Nx = max(abs(Nx1), abs(Nx2))
            max_sigma = max(sigma1, sigma2)
            if max_Nx > max_force:
                max_force = max_Nx
                max_force_elem = res['index']
            if max_sigma > max_stress:
                max_stress = max_sigma
                max_stress_elem = res['index']
            self.analysis_text.insert(tk.END, f"   Стержень {res['index']} ({res['nodes'][0]}-{res['nodes'][1]}):\n")
            self.analysis_text.insert(tk.END, f"      Nx1 = {Nx1:.2f} Н, Nx2 = {Nx2:.2f} Н\n")
            self.analysis_text.insert(tk.END, f"      σx1 = {sigma1:.2f} Па, σx2 = {sigma2:.2f} Па\n")
            self.analysis_text.insert(tk.END, f"      Допускаемое: {res['allowable_sigma']:.2f} Па\n")
        if max_force_elem:
            self.analysis_text.insert(tk.END, f"   Максимальное усилие: стержень {max_force_elem}, |Nx| = {max_force:.2f} Н\n")
        if max_stress_elem:
            self.analysis_text.insert(tk.END, f"   Максимальное напряжение: стержень {max_stress_elem}, |σx| = {max_stress:.2f} Па\n")
        self.analysis_text.insert(tk.END, "\n")
        
        self.analysis_text.insert(tk.END, "4. ПРОВЕРКА ДОПУСКАЕМЫХ НАПРЯЖЕНИЙ:\n")
        warnings = self.results.get('warnings', [])
        if warnings:
            for w in warnings:
                self.analysis_text.insert(tk.END, f"    {w}\n")
        else:
            self.analysis_text.insert(tk.END, "   Все напряжения в пределах допуска.\n")
        self.analysis_text.insert(tk.END, "\n")
        
        self.analysis_text.insert(tk.END, "")
        self.analysis_text.configure(state='disabled')

    def run_analysis(self):
        try:
            processor = StructuralProcessor(self.nodes, self.elements, self.supports, self.forces, self.element_forces)
            self.results = processor.solve()
        except ValueError as exc:
            messagebox.showerror("Расчёт", str(exc))
            return
        self.populate_results_ui()
        self.section_result_var.set("Расчёт выполнен. Выберите стержень и x.")
        messagebox.showinfo("Расчёт", "Расчёт завершён.")
        self.draw_system()

    def calculate_section_values(self):
        if not self.results:
            messagebox.showwarning("Сечение", "Сначала выполните расчёт.")
            return
        combo_idx = self.section_element_combo.current()
        if combo_idx == -1 or not getattr(self, 'sorted_elements_for_combo', []):
            messagebox.showerror("Сечение", "Выберите стержень.")
            return
        try:
            element_idx, _ = self.sorted_elements_for_combo[combo_idx]
            xi = float(self.entry_section_coordinate.get())
        except Exception:
            messagebox.showerror("Сечение", "Введите корректное значение x.")
            return
        element_result = self.results['element_results'][element_idx]
        # Дополнительная проверка диапазона с указанием конкретной длины L выбранного стержня
        L = element_result['length']
        if not (0.0 <= xi <= L):
            messagebox.showerror(
                "Сечение",
                f"Локальная координата должна быть в пределах [0, {L:.3f}] м для выбранного стержня."
            )
            return
        try:
            values = StructuralProcessor.evaluate_section(element_result, xi)
        except ValueError as exc:
            messagebox.showerror("Сечение", str(exc))
            return
        self.section_result_var.set(f"ux={values['ux']:.5f}; Nx={values['Nx']:.2f}; σx={values['sigma']:.2f}")

    def show_diagram(self, component):
        if not self.results:
            messagebox.showwarning("Диаграмма", "Сначала выполните расчёт.")
            return
        # Проверяем корректность введённого числа точек
        try:
            self._get_diagram_points_per_element(strict=True)
        except ValueError:
            return

        segments = self._build_diagram_segments(component)
        if not segments:
            messagebox.showwarning("Диаграмма", "Нет данных для построения.")
            return
        titles = {
            'ux': ("Эпюра ux", "ux"),
            'Nx': ("Эпюра Nx", "Nx"),
            'sigma': ("Эпюра σx", "σx"),
        }
        title, ylabel = titles[component]
        DiagramWindow(self.root, segments, title, ylabel, component)

    def _get_diagram_points_per_element(self, strict: bool = False):
        """
        Возвращает число точек на стержень, заданное пользователем.
        strict=True — при некорректном вводе показывает ошибку и выбрасывает ValueError.
        strict=False — тихо подставляет разумные значения (по умолчанию 20, минимум 2).
        """
        default_points = 20
        if not hasattr(self, "entry_diagram_points"):
            return default_points
        text = self.entry_diagram_points.get().strip()
        if not text:
            if strict:
                messagebox.showerror("Диаграммы", "Число точек на стержень должно быть целым числом ≥ 2.")
                raise ValueError
            return default_points
        try:
            value = int(text)
        except ValueError:
            if strict:
                messagebox.showerror("Диаграммы", "Число точек на стержень должно быть целым числом ≥ 2.")
                raise ValueError
            return default_points

        if value < 2:
            if strict:
                messagebox.showerror("Диаграммы", "Число точек на стержень должно быть не меньше 2.")
                raise ValueError
            return 2
        return value

    def _build_diagram_segments(self, component):
        if not self.results:
            return []
        if getattr(self, 'sorted_elements_for_combo', None):
            sorted_elements = self.sorted_elements_for_combo
        else:
            sorted_elements = list(enumerate(self.elements))

        segments = []
        num_points_plot = 20
        num_points_table = self._get_diagram_points_per_element()

        for idx, _ in sorted_elements:
            if idx >= len(self.results['element_results']):
                continue
            res = self.results['element_results'][idx]
            length = res['length']
            q_val = res['q']

            plot_points = []
            for i in range(num_points_plot):
                x_local_plot = length * i / (num_points_plot - 1)
                values_plot = StructuralProcessor.evaluate_section(res, x_local_plot)
                if component == 'ux':
                    plot_points.append((x_local_plot, values_plot['ux']))
                elif component == 'Nx':
                    plot_points.append((x_local_plot, values_plot['Nx']))
                else:  # 'sigma'
                    plot_points.append((x_local_plot, values_plot['sigma']))

            values_list = []
            for i in range(num_points_table):
                x_local_table = length * i / (num_points_table - 1)
                values = StructuralProcessor.evaluate_section(res, x_local_table)
                values_list.append({
                    'x': x_local_table,
                    'ux': values['ux'],
                    'Nx': values['Nx'],
                    'sigma': values['sigma'],
                })

            segments.append({
                'index': res['index'],
                'length': length,
                'points': plot_points,
                'values': values_list,
            })
        return segments

    def export_report(self):
        if not self.results:
            messagebox.showwarning("Отчёт", "Сначала выполните расчёт.")
            return
        report = {
            'timestamp': datetime.now().isoformat(sep=' ', timespec='seconds'),
            'nodes': self.nodes,
            'elements': self.elements,
            'supports': self.supports,
            'forces': self.forces,
            'element_forces': self.element_forces,
            'results': self.results,
        }
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("Word","*.docx"), ("JSON","*.json"), ("Текст","*.txt")]
        )
        if not file_path:
            return
        try:
            if file_path.lower().endswith(".txt"):
                content = self._format_report_text(report)
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)
            elif file_path.lower().endswith(".docx"):
                self._export_docx_report(file_path, report)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(report, f, ensure_ascii=False, indent=2)
        except OSError as exc:
            messagebox.showerror("Отчёт", f"Не удалось сохранить файл: {exc}")
            return
        except RuntimeError as exc:
            messagebox.showerror("Отчёт", str(exc))
            return
        messagebox.showinfo("Отчёт", "Файл сохранён.")

    def _format_report_text(self, report):
        lines = [
            "Отчёт по расчёту плоской стержневой системы",
            f"Дата: {report['timestamp']}",
            "",
            "Узлы:",
        ]
        for i, x in enumerate(report['nodes'], start=1):
            lines.append(f"  {i}: x = {x:.3f}")
        lines.append("\nСтержни:")
        for idx, element in enumerate(report['elements'], start=1):
            n1, n2 = element['nodes']
            lines.append(
                f"  {idx}: {n1}-{n2}, L={element['L']:.3f}, A={element['A']}, E={element['E']}, σдоп={element['sigma']}"
            )
        lines.append("\nОпоры: " + (", ".join(map(str, report['supports'])) if report['supports'] else "нет"))
        if report['forces']:
            lines.append("Сосредоточенные силы:")
            for node, value in report['forces']:
                lines.append(f"  узел {node}: F = {value}")
        if report['element_forces']:
            lines.append("Распределённые q:")
            for idx, q in report['element_forces']:
                n1, n2 = report['elements'][idx]['nodes']
                lines.append(f"  стержень {idx+1} ({n1}-{n2}): q = {q}")
        lines.append("\nПеремещения узлов:")
        for i, ux in enumerate(report['results']['displacements'], start=1):
            lines.append(f"  узел {i}: ux = {ux:.6f}")
        lines.append("\nРезультаты по стержням:")
        for res in report['results']['element_results']:
            lines.append(
                f"  {res['index']}: {res['nodes'][0]}-{res['nodes'][1]}, Nx={res['force']:.2f}, σx={res['stress']:.2f}"
            )
        lines.append("\nРеакции опор:")
        for node, value in report['results'].get('reactions', []):
            lines.append(f"  R{node} = {value:.2f}")
        warnings = report['results'].get('warnings', [])
        if warnings:
            lines.append("\nПредупреждения:")
            lines.extend([f"  - {w}" for w in warnings])
        else:
            lines.append("\nПредупреждений нет.")
        return "\n".join(lines)

    def _export_docx_report(self, filepath, report):
        if Document is None or Inches is None:
            raise RuntimeError("Библиотека python-docx недоступна. Установите её командой 'pip install python-docx'.")

        doc = Document()
        doc.add_heading("Отчёт по расчёту плоской стержневой системы", level=1)
        doc.add_paragraph(f"Дата: {report['timestamp']}")

        self._add_docx_table(
            doc,
            "Узлы",
            ["Узел", "X, м"],
            [[str(i+1), f"{x:.3f}"] for i, x in enumerate(report['nodes'])]
        )

        element_rows = []
        for idx, element in enumerate(report['elements'], start=1):
            n1, n2 = element['nodes']
            element_rows.append([
                str(idx),
                f"{n1}-{n2}",
                f"{element['L']:.3f}",
                f"{element['A']}",
                f"{element['E']}",
                f"{element['sigma']}"
            ])
        self._add_docx_table(
            doc,
            "Стержни",
            ["#", "Узлы", "L, м", "A, м²", "E, Па", "σдоп, Па"],
            element_rows
        )

        if report['forces']:
            self._add_docx_table(
                doc,
                "Силы",
                ["Узел", "F, Н"],
                [[str(n), f"{value}"] for n, value in report['forces']]
            )
        if report['element_forces']:
            q_rows = []
            for idx, q in report['element_forces']:
                n1, n2 = report['elements'][idx]['nodes']
                q_rows.append([f"{idx+1} ({n1}-{n2})", f"{q}"])
            self._add_docx_table(
                doc,
                "Распределённые нагрузки q",
                ["Стержень", "q, Н/м"],
                q_rows
            )

        self._add_docx_table(
            doc,
            "Перемещения",
            ["Узел", "ux, м"],
            [[str(i+1), f"{ux:.6f}"] for i, ux in enumerate(report['results']['displacements'])]
        )

        element_res_rows = []
        for res in report['results']['element_results']:
            length = res['length']
            q_val = res['q']

            Nx1 = res['force'] + q_val * length / 2
            Nx2 = res['force'] - q_val * length / 2

            sigma1 = Nx1 / res['A']
            sigma2 = Nx2 / res['A']

            element_res_rows.append([
                f"{res['index']}",
                f"{res['nodes'][0]}-{res['nodes'][1]}",
                f"{Nx1:.2f}",
                f"{Nx2:.2f}",
                f"{res['u1']:.6f}",
                f"{res['u2']:.6f}",
                f"{sigma1:.2f}",
                f"{sigma2:.2f}",
            ])
        self._add_docx_table(
            doc,
            "Результаты по стержням",
            ["#", "Узлы", "Nx1, Н", "Nx2, Н", "ux1, м", "ux2, м", "σx1, Па", "σx2, Па"],
            element_res_rows
        )

        reactions = report['results'].get('reactions', [])
        if reactions:
            self._add_docx_table(
                doc,
                "Реакции опор",
                ["Узел", "R, Н"],
                [[str(node), f"{val:.2f}"] for node, val in reactions]
            )

        warnings = report['results'].get('warnings', [])
        if warnings:
            doc.add_heading("Предупреждения", level=2)
            for warning in warnings:
                doc.add_paragraph(warning, style='List Bullet')
            doc.add_paragraph("Обнаружены превышения допускаемых напряжений. См. предупреждения выше.")
        else:
            doc.add_heading("Анализ напряжений", level=2)
            doc.add_paragraph("Все напряжения в пределах допуска.")

        with tempfile.TemporaryDirectory(prefix="sapr_report_") as temp_dir:
            if plt and self.results:
                title_map = {'ux': "График ux", 'Nx': "Эпюра Nx", 'sigma': "Эпюра σx"}
                for comp in ('ux', 'Nx', 'sigma'):
                    img_path = os.path.join(temp_dir, f"{comp}.png")
                    if self._save_diagram_png(comp, img_path):
                        doc.add_heading(title_map[comp], level=2)
                        doc.add_picture(img_path, width=Inches(5.5))
                        doc.add_paragraph(self._diagram_caption(comp))

        # Принудительно делаем весь текст чёрным
        if 'RGBColor' in globals():
            from docx.enum.style import WD_STYLE_TYPE
            # Заголовки и базовые стили
            for style in doc.styles:
                if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
                    if style.font is not None and style.font.color is not None:
                        style.font.color.rgb = RGBColor(0, 0, 0)

            # Все параграфы и таблицы
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)

        doc.save(filepath)

    @staticmethod
    def _add_docx_table(doc, title, headers, rows):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            if 'RGBColor' in globals():
                run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = value

    def _diagram_caption(self, component):
        captions = {
            'ux': "Перемещения узлов вдоль оси стержневой системы.",
            'Nx': "Продольные силы Nx по длине конструкции.",
            'sigma': "Нормальные напряжения σx по длине конструкции."
        }
        return captions.get(component, "")

    def _save_diagram_png(self, component, file_path):
        if not plt:
            return False
        series = self._diagram_component_series(component)
        if not series or len(series[0]) < 2:
            return False

        x_values, y_values = series
        fig, ax = plt.subplots(figsize=(6, 3))
        ax.plot(x_values, y_values, color="#1f77b4", linewidth=2)
        ax.fill_between(x_values, y_values, color="#aec7e8", alpha=0.4)
        ax.axhline(0, color="#444444", linewidth=1, linestyle="--")
        labels = {'ux': "ux, м", 'Nx': "Nx, Н", 'sigma': "σx, Па"}
        ax.set_xlabel("Глобальная координата, м")
        ax.set_ylabel(labels.get(component, ""))
        ax.grid(True, alpha=0.3)
        ax.set_title(self._diagram_caption(component))
        fig.tight_layout()
        fig.savefig(file_path, dpi=200)
        plt.close(fig)
        return True

    def _diagram_component_series(self, component):
        if not self.nodes or not self.elements or not self.results:
            return None
        segments = self._build_diagram_segments(component)
        if not segments:
            return None

        x_values = []
        y_values = []
        global_offset = 0.0
        for seg in segments:
            for x_local, value in seg['points']:
                x_values.append(global_offset + x_local)
                y_values.append(value)
            global_offset += seg['length']

        return x_values, y_values


    #Добавление
    def add_node(self):
        try: 
            x = float(self.entry_node.get())
        except: 
            messagebox.showerror("Ошибка","Введите число для X")
            return
        if x in self.nodes: 
            messagebox.showerror("Ошибка","Узел с такой координатой уже существует!") 
            return
        self.nodes.append(x)
        self.nodes.sort()
        self.reset_view_range()
        self.reset_postprocessor()
        self.update_comboboxes()
        self.entry_node.delete(0,tk.END)
        self.refresh_lists()
        self.draw_system()

    def add_element(self):
        try:
            n1 = int(self.combo_n1.get())
            n2 = int(self.combo_n2.get())
        except:
            messagebox.showerror("Ошибка","Выберите начальный и конечный узел") 
            return

        try:
            A = float(self.entry_A.get().strip()) if self.entry_A.get().strip() else DEFAULT_A
            if A <= 0: 
                messagebox.showerror("Ошибка","Площадь A должна быть больше 0")
                return
        except:
            messagebox.showerror("Ошибка","Некорректное значение A")
            return

        try:
            E = float(self.entry_E.get().strip()) if self.entry_E.get().strip() else DEFAULT_E
            if E <= 0: 
                messagebox.showerror("Ошибка","Модуль упругости E должен быть больше 0")
                return
        except:
            messagebox.showerror("Ошибка","Некорректное значение E")
            return

        try:
            sigma = float(self.entry_sigma.get().strip()) if self.entry_sigma.get().strip() else DEFAULT_SIGMA
            if sigma <= 0: 
                messagebox.showerror("Ошибка","Допускаемое напряжение σ должно быть больше 0")
                return
        except:
            messagebox.showerror("Ошибка","Некорректное значение σ")
            return

        if n1 == n2:
            messagebox.showerror("Ошибка","Начальный и конечный узлы не могут совпадать")
            return

        x1_new = min(self.nodes[n1-1], self.nodes[n2-1])
        x2_new = max(self.nodes[n1-1], self.nodes[n2-1])

        for e in self.elements:
            x1_e = min(self.nodes[e['nodes'][0]-1], self.nodes[e['nodes'][1]-1])
            x2_e = max(self.nodes[e['nodes'][0]-1], self.nodes[e['nodes'][1]-1])
            if max(x1_e, x1_new) < min(x2_e, x2_new):
                messagebox.showerror("Ошибка","Новый стержень пересекается с существующим!")
                return

        for i, x in enumerate(self.nodes, start=1):
            if i != n1 and i != n2:
                if x1_new < x < x2_new:
                    messagebox.showerror("Ошибка", f"Стержень не может пересекать узел {i}")
                    return

        L = abs(self.nodes[n2-1] - self.nodes[n1-1])
        self.elements.append({'nodes': (n1, n2), 'A': A, 'E': E, 'sigma': sigma, 'L': L})

        self.reset_postprocessor()
        self.entry_A.delete(0, tk.END)
        self.entry_E.delete(0, tk.END)
        self.entry_sigma.delete(0, tk.END)

        self.combo_n1.set('')
        self.combo_n2.set('')

        self.update_comboboxes()
        self.refresh_lists()
        self.draw_system()

    def add_q(self):
        try:
            combo_idx = self.combo_q_element.current()
            if combo_idx == -1:
                raise ValueError
            idx, _ = self.sorted_elements_for_combo[combo_idx]
            q_val = float(self.entry_q.get())
        except:
            messagebox.showerror("Ошибка","Некорректные данные")
            self.entry_q.delete(0, tk.END)
            self.combo_q_element.set('')
            return

        if q_val == 0:
            messagebox.showerror("Ошибка","Нельзя добавлять распределённую нагрузку q=0")
            self.entry_q.delete(0, tk.END)
            self.combo_q_element.set('')
            return

        if any(e_idx == idx for e_idx, _ in self.element_forces):
            messagebox.showerror("Ошибка","На этом стержне уже есть распределённая нагрузка q")
            self.entry_q.delete(0, tk.END)
            self.combo_q_element.set('')
            return

        self.element_forces.append((idx, q_val))
        self.reset_postprocessor()
        self.entry_q.delete(0, tk.END)
        self.combo_q_element.set('')
        self.refresh_lists()
        self.draw_system()

    def add_support(self):
        try:
            n = int(self.combo_support.get())
        except:
            messagebox.showerror("Ошибка","Выберите узел для опоры")
            return

        if n in self.supports:
            messagebox.showerror("Ошибка", "На этом узле уже есть опора")
            return

        self.supports.append(n)
        self.reset_postprocessor()
        self.combo_support.set('')
        self.refresh_lists()
        self.draw_system()

    def add_force(self):
        try:
            n = int(self.combo_force_node.get())
            F = float(self.entry_F.get())
        except:
            messagebox.showerror("Ошибка","Некорректные данные")
            self.entry_F.delete(0, tk.END)
            self.combo_force_node.set('')
            return

        if F == 0:
            messagebox.showerror("Ошибка", "Сила F не может быть равна 0")
            self.entry_F.delete(0, tk.END)
            self.combo_force_node.set('')
            return

        if any(f[0] == n for f in self.forces):
            messagebox.showerror("Ошибка","На узле уже есть сила")
            self.entry_F.delete(0, tk.END)
            self.combo_force_node.set('')
            return

        self.forces.append((n,F))
        self.reset_postprocessor()
        self.entry_F.delete(0, tk.END)
        self.combo_force_node.set('')
        self.refresh_lists()
        self.draw_system()

    #Удаление
    def delete_node(self):
        selected = self.tree_nodes.selection()
        if not selected:
            messagebox.showwarning("Удаление", "Выберите узел")
            return

        tree_idx = self.tree_nodes.index(selected[0])
        sorted_nodes = sorted(self.nodes)
        node_to_remove = sorted_nodes[tree_idx]
        for e in self.elements:
            n1, n2 = e['nodes']
            if self.nodes[n1-1] == node_to_remove or self.nodes[n2-1] == node_to_remove:
                messagebox.showerror("Ошибка", "Нельзя удалить узел, входящий в стержень")
                return

        self.nodes.remove(node_to_remove)
        self.reset_view_range()
        self.reset_postprocessor()
        self.update_comboboxes()
        self.refresh_lists()
        self.draw_system()

    def delete_element(self):
        selected = self.tree_elements.selection()
        if not selected:
            messagebox.showwarning("Удаление","Выберите стержень")
            return
        idx = self.tree_elements.index(selected[0])
        sorted_elements = sorted(enumerate(self.elements), key=lambda x: min(self.nodes[x[1]['nodes'][0]-1], self.nodes[x[1]['nodes'][1]-1]))
        real_idx, _ = sorted_elements[idx]
        del self.elements[real_idx]
        self.element_forces = [(i-(1 if i>real_idx else 0), q) for i,q in self.element_forces if i!=real_idx]
        self.reset_postprocessor()
        self.update_comboboxes()
        self.refresh_lists()
        self.draw_system()

    def delete_q(self):
        selected = self.tree_q.selection()
        if not selected:
            messagebox.showwarning("Удаление","Выберите нагрузку q")
            return
        idx = self.tree_q.index(selected[0])
        sorted_q = sorted(self.element_forces, key=lambda x: min(self.nodes[self.elements[x[0]]['nodes'][0]-1], self.nodes[self.elements[x[0]]['nodes'][1]-1]))
        real_idx = self.element_forces.index(sorted_q[idx])
        del self.element_forces[real_idx]
        self.reset_postprocessor()
        self.refresh_lists()
        self.draw_system()

    def delete_support(self):
        selected = self.tree_supports.selection()
        if not selected:
            messagebox.showwarning("Удаление","Выберите опору")
            return

        tree_idx = self.tree_supports.index(selected[0])
        sorted_supports = sorted(self.supports)
        support_to_remove = sorted_supports[tree_idx]

        self.supports.remove(support_to_remove)
        self.reset_postprocessor()
        self.refresh_lists()
        self.draw_system()

    def delete_force(self):
        selected = self.tree_forces.selection()
        if not selected:
            messagebox.showwarning("Удаление","Выберите силу")
            return

        tree_idx = self.tree_forces.index(selected[0])
        sorted_forces = sorted(self.forces, key=lambda x: x[0])
        force_to_remove = sorted_forces[tree_idx]

        self.forces = [f for f in self.forces if f != force_to_remove]

        self.reset_postprocessor()
        self.refresh_lists()
        self.draw_system()


    #Очистить все
    def clear_all(self):
        if messagebox.askyesno("Подтверждение","Очистить всё?"):
            self.nodes.clear()
            self.elements.clear()
            self.supports.clear()
            self.forces.clear()
            self.element_forces.clear()
            self.imported_file_name = None
            if hasattr(self, 'file_name_label'):
                self.file_name_label.config(text="")

            self.combo_n1.set('')
            self.combo_n2.set('')
            self.combo_q_element.set('')
            self.combo_support.set('')
            self.combo_force_node.set('')

            self.reset_view_range()
            self.reset_postprocessor()
            self.update_comboboxes()
            self.refresh_lists()
            self.canvas.delete("all")


    #Масштаб
    def scale_canvas(self, factor, anchor_pixel=None):
        if not self.nodes or factor <= 0:
            return
        if self.view_min is None or self.view_max is None:
            self.reset_view_range()
        if self.view_min is None or self.view_max is None:
            return

        data_min = min(self.nodes)
        data_max = max(self.nodes)
        span = self.view_max - self.view_min
        if math.isclose(span, 0.0):
            span = 1.0

        if anchor_pixel is None:
            anchor_world = self.view_min + span / 2
        else:
            anchor_world = self._pixel_to_world(anchor_pixel)

        new_min = anchor_world - (anchor_world - self.view_min) / factor
        new_max = anchor_world + (self.view_max - anchor_world) / factor

        data_span = max(data_max - data_min, 1e-6)
        new_span = new_max - new_min
        if new_span > data_span:
            new_min, new_max = data_min, data_max
        else:
            if new_min < data_min:
                shift = data_min - new_min
                new_min += shift
                new_max += shift
            if new_max > data_max:
                shift = new_max - data_max
                new_min -= shift
                new_max -= shift

        if math.isclose(new_max, new_min):
            padding = 0.5
            new_min -= padding
            new_max += padding

        self.view_min = new_min
        self.view_max = new_max
        self.draw_system()

    def reset_scale(self):
        self.reset_view_range()
        self.draw_system()

    def reset_view_range(self):
        if not self.nodes:
            self.view_min = None
            self.view_max = None
            return
        self.view_min = min(self.nodes)
        self.view_max = max(self.nodes)
        if math.isclose(self.view_min, self.view_max):
            self.view_max = self.view_min + 1.0

    def _pixel_to_world(self, pixel_x):
        if self.view_min is None or self.view_max is None:
            return 0.0
        w = max(self.canvas.winfo_width(), 1)
        margin = 80
        usable = max(w - 2 * margin, 1)
        rel = (pixel_x - margin) / usable
        rel = min(max(rel, 0.0), 1.0)
        return self.view_min + rel * (self.view_max - self.view_min)

    def on_mousewheel(self,event):
        anchor = getattr(event, "x", None)
        factor = 1.1 if event.delta>0 else 0.9
        self.scale_canvas(factor, anchor)

    #Рисуем
    def draw_system(self):
        self.canvas.delete("all")
        if not self.nodes: 
            return
        if self.view_min is None or self.view_max is None:
            self.reset_view_range()
        if self.view_min is None or self.view_max is None:
            return
        w,h = self.canvas.winfo_width(), self.canvas.winfo_height()
        margin = 80
        x_min, x_max = self.view_min, self.view_max
        scale = (w-2*margin)/(x_max-x_min+1e-9)
        y_mid = h/2

        if self.nodes:
            for node_x in self.nodes:
                x_grid = margin + (node_x - x_min) * scale
                if margin <= x_grid <= w - margin:
                    self.canvas.create_line(x_grid, 0, x_grid, h, fill="#f0f0f0", width=1)

        num_y_grid = 15
        for i in range(num_y_grid + 1):
            y_grid = h * i / num_y_grid
            self.canvas.create_line(margin, y_grid, w - margin, y_grid, fill="#f0f0f0", width=1)

        min_width = 15.0
        max_width = 50.0
        if self.elements:
            areas = [e['A'] for e in self.elements]
            min_area = min(areas)
            max_area = max(areas)
        else:
            min_area = max_area = 1.0

        rod_rects = []
        element_label_map = {}
        if self.elements:
            sorted_indices = sorted(
                range(len(self.elements)),
                key=lambda idx: min(
                    self.nodes[self.elements[idx]['nodes'][0]-1],
                    self.nodes[self.elements[idx]['nodes'][1]-1]
                )
            )
            element_label_map = {idx: order + 1 for order, idx in enumerate(sorted_indices)}

        max_bottom = y_mid 
        for idx,e in enumerate(self.elements):
            n1,n2 = e['nodes']
            x1 = margin + (self.nodes[n1-1]-x_min)*scale
            x2 = margin + (self.nodes[n2-1]-x_min)*scale

            if math.isclose(max_area, min_area):

                rod_height = (min_width + max_width) / 2
            else:
                area_ratio = (e['A'] - min_area) / (max_area - min_area)
                rod_height = min_width + area_ratio * (max_width - min_width)

            rect_top = y_mid - rod_height / 2
            rect_bottom = y_mid + rod_height / 2
            max_bottom = max(max_bottom, rect_bottom)
            self.canvas.create_rectangle(x1, rect_top, x2, rect_bottom, outline="black", width=2, fill="")
            label_value = element_label_map.get(idx, idx + 1)
            self.canvas.create_text((x1+x2)/2, y_mid-rod_height/2-15, text=f"{label_value}", fill="black", font=("Arial",9))
            rod_rects.append((idx, x1, x2, rect_top, rect_bottom, y_mid))

        max_q_y = max_bottom
        for idx,q_val in self.element_forces:
            e = self.elements[idx]
            x1 = margin + (self.nodes[e['nodes'][0]-1]-x_min)*scale
            x2 = margin + (self.nodes[e['nodes'][1]-1]-x_min)*scale

            rod_height = max_width  
            for rod_idx, rx1, rx2, rtop, rbottom, rymid in rod_rects:
                if rod_idx == idx:
                    rod_height = rbottom - rtop
                    break
            q_y = y_mid + rod_height / 2 + 15  
            max_q_y = max(max_q_y, q_y + 25) 
            num_arrows = max(2, int(abs(x2-x1)/20))
            step = (x2-x1)/num_arrows
            for i_arrow in range(num_arrows):
                px = x1 + i_arrow*step
                if q_val>0: self.canvas.create_line(px,q_y,px+10,q_y,fill="blue",arrow=tk.LAST,width=2)
                else: self.canvas.create_line(px,q_y,px-10,q_y,fill="blue",arrow=tk.LAST,width=2)
            self.canvas.create_text((x1+x2)/2, q_y + 15, text=f"q={q_val}", fill="blue", font=("Arial",9))
        
        axis_y = max_q_y + 20
        self.canvas.create_line(margin, axis_y, w-margin, axis_y, width=2, arrow=tk.LAST)

        for i,x in enumerate(self.nodes,start=1):
            X = margin + (x-x_min)*scale
            self.canvas.create_oval(X-3,y_mid-3,X+3,y_mid+3,fill="black")
            self.canvas.create_text(X, axis_y + 15, text=f"{i}", font=("Arial",9,"bold"))
            self.canvas.create_text(X, axis_y + 30, text=f"{x:.2f}",font=("Arial",8))

        system_center_x = (x_min + x_max) / 2
        system_center_pixel = margin + (system_center_x - x_min) * scale
        
        for n in self.supports:
            X = margin + (self.nodes[n-1]-x_min)*scale
            rod_height = max_width  
            for rod_idx, rx1, rx2, rtop, rbottom, rymid in rod_rects:

                if min(rx1, rx2) <= X <= max(rx1, rx2):
                    rod_height = rbottom - rtop
                    break

            support_top = y_mid - rod_height / 2
            support_bottom = y_mid + rod_height / 2
            self.canvas.create_line(X, support_top, X, support_bottom, fill="red", width=3)

            is_left = X < system_center_pixel
            
            hatch_length = 8
            hatch_spacing = 6
            num_hatches = int((support_bottom - support_top) / hatch_spacing)
            for i in range(num_hatches):
                y_pos = support_top + i * hatch_spacing
                if is_left:
                    x1_hatch = X - hatch_length
                    y1_hatch = y_pos + hatch_length
                    x2_hatch = X
                    y2_hatch = y_pos
                else:
                    x1_hatch = X
                    y1_hatch = y_pos
                    x2_hatch = X + hatch_length
                    y2_hatch = y_pos + hatch_length
                self.canvas.create_line(x1_hatch, y1_hatch, x2_hatch, y2_hatch, fill="red", width=2)

        for n,F in self.forces:
            X = margin + (self.nodes[n-1]-x_min)*scale

            found_rod = False
            for rod_idx, x1, x2, rect_top, rect_bottom, rod_y_mid in rod_rects:
                if min(x1, x2) <= X <= max(x1, x2):

                    arrow_len = min(25, (max(x1, x2) - min(x1, x2)) * 0.3)
                    if F > 0:
                        arrow_end_x = X + arrow_len
                        self.canvas.create_line(X, rod_y_mid, arrow_end_x, rod_y_mid, fill="green", arrow=tk.LAST, width=2)

                        self.canvas.create_text(arrow_end_x + 15, rod_y_mid, text=f"F={F}", fill="green", font=("Arial", 8, "bold"), anchor=tk.W)
                    else:
                        arrow_end_x = X - arrow_len
                        self.canvas.create_line(X, rod_y_mid, arrow_end_x, rod_y_mid, fill="green", arrow=tk.LAST, width=2)

                        self.canvas.create_text(arrow_end_x - 15, rod_y_mid, text=f"F={F}", fill="green", font=("Arial", 8, "bold"), anchor=tk.E)
                    found_rod = True
                    break

            if not found_rod:
                arrow_len = 20
                if F > 0:
                    self.canvas.create_line(X, y_mid, X + arrow_len, y_mid, fill="green", arrow=tk.LAST, width=2)
                else:
                    self.canvas.create_line(X, y_mid, X - arrow_len, y_mid, fill="green", arrow=tk.LAST, width=2)
                self.canvas.create_text(X, y_mid - 15, text=f"F={F}", fill="green", font=("Arial", 9))

        
            
        bbox = self.canvas.bbox("all")
        if bbox:
            self.canvas.configure(scrollregion=bbox)
        else:
            self.canvas.configure(scrollregion=(0, 0, w, h))
    # Экспорт
    def export_json(self):
        data = {'nodes':self.nodes,'elements':self.elements,'supports':self.supports,'forces':self.forces,'element_forces':self.element_forces}
        file_path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON","*.json")])
        if file_path:
            with open(file_path,"w") as f: json.dump(data,f)
            messagebox.showinfo("Экспорт","Проект сохранён.")

    #Импорт
    def import_json(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON","*.json")])
        if file_path:
            with open(file_path,"r") as f:
                data = json.load(f)

            self.nodes = data.get('nodes', [])
            self.elements = data.get('elements', [])
            self.supports = data.get('supports', [])
            self.forces = data.get('forces', [])

            self.element_forces.clear()
            for idx, q in data.get('element_forces', []):
                if 0 <= idx < len(self.elements) and q != 0:
                    self.element_forces.append((idx, q))

            self.imported_file_name = os.path.basename(file_path)
            if hasattr(self, 'file_name_label'):
                self.file_name_label.config(text=f"Файл: {self.imported_file_name}")

            self.reset_view_range()
            self.reset_postprocessor()
            self.update_comboboxes()
            self.refresh_lists()
            self.draw_system()
            messagebox.showinfo("Импорт","Проект загружен.")


if __name__=="__main__":
    root = tk.Tk()
    root.iconbitmap("icon.ico")
    app = PreprocessorApp(root)
    root.mainloop()
